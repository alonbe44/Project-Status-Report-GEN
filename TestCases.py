import os
import smtplib
import unittest
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from main import send_email, Gen_Document

# Author: Abedalrahman Rasem


"""
Code Analysis

Objective:
The objective of the 'send_email' function is to send an email with an attachment. The function takes the filename of the attachment as input and sends an email to the specified recipient with the attachment.

Inputs:
The 'send_email' function takes a single input, which is the filename of the attachment to be sent.

Flow:
The 'send_email' function first creates a message object and sets the email details such as sender, recipient, and subject. It then attaches the specified file to the email and connects to the email server. The function starts a secure TLS connection, logs in to the email account, sends the email, and quits the email server connection.

Outputs:
The main output of the 'send_email' function is an email sent to the specified recipient with the specified attachment.

Additional aspects:
- The function checks if the specified attachment file exists before attaching it to the email.
- The function raises a 'FileNotFoundError' exception if the specified attachment file is not found.
- The function handles any other exceptions that may occur during the email sending process and raises them with additional information.
"""
class TestSendEmail(unittest.TestCase):
    #  Tests that an email is sent successfully
    def test_successful_email(self,mocker):
        self.mocker=mocker
        try:
            send_email('test.txt')
        except:
            self.fail('An exception was raised')
        self.assertTrue(True)

    #  Tests that an error is raised when attachment file is not found
    def test_attachment_not_found(self):
        with self.assertRaises(FileNotFoundError):
            send_email('nonexistent_file.txt')

    #  Tests that an error is raised when there is an SMTP login error
    def test_smtp_login_error(self):
        with self.assertRaises(smtplib.SMTPAuthenticationError):
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login('globtesting0@gmail.com', 'wrong_password')
            send_email('test_attachment.txt')
            server.quit()

    #  Tests that email details (sender, recipient, subject) are set correctly
    def test_email_details(self):
        msg=send_email('test.txt')
        self.assertEqual(msg['From'], 'globtesting0@gmail.com')
        self.assertEqual(msg['To'], 'Abdelrahman.Rasem@globitel.com')
        self.assertEqual(msg['Subject'], 'Project Status ' + str(datetime.now().isocalendar()[1]) + ' 2023')




    def test_successful_email(self, mocker):
        mocker.patch('smtplib.SMTP')
        mocker.patch('smtplib.SMTP.starttls')
        mocker.patch('smtplib.SMTP.login')
        mocker.patch('smtplib.SMTP.send_message')
        mocker.patch('smtplib.SMTP.quit')
        send_email('test.txt')
        smtplib.SMTP.assert_called_once_with('smtp.gmail.com', 587)
        smtplib.SMTP.starttls.assert_called_once()
        smtplib.SMTP.login.assert_called_once_with('globtesting0@gmail.com', 'vfqmlzzwiubkyxlg')
        smtplib.SMTP.send_message.assert_called_once()
        smtplib.SMTP.quit.assert_called_once()


"""
Code Analysis

Objective:
The Gen_Document function aims to generate a Word document containing a table with project status information. The function takes a hash map as input, which contains project details such as project name, engineer name, related to, status, and comments. The function highlights rows based on the status value and saves the document with a filename that includes the current week number.

Inputs:
- mp: a hash map containing project details such as project name, engineer name, related to, status, and comments.

Flow:
1. Create a new Word document.
2. Add a header to the document.
3. Add the current date on the left side of the header.
4. Create a table with headers.
5. Iterate over the hash map and add data to the table.
6. Highlight rows based on the status value.
7. Save the document with a filename that includes the current week number.
8. Open the Word document.
9. Modify the page size for landscape orientation.
10. Adjust the width of the last column in the table.
11. Save the modified document.

Outputs:
- A Word document containing a table with project status information.
- The filename of the saved document.

Additional aspects:
- The function uses the docx library to create and modify Word documents.
- The function highlights rows based on the status value using different colors.
- The function adjusts the page size and column width for better document formatting.
- The function saves the document with a filename that includes the current week number.
- The function opens the saved document for further modification if needed.
"""
class TestGenDocument(unittest.TestCase):
    #  Tests that a Word document is generated with a header and table
    def test_generate_document(self):
        mp = {
            'project name': {1: 'Project 1', 2: 'Project 2'},
            'Engineer name': {1: 'Engineer 1', 2: 'Engineer 2'},
            'Related To': {1: 'Related To 1', 2: 'Related To 2'},
            'status': {1: 'In Progress', 2: 'Completed'},
            'Details': {1: 'Details 1', 2: 'Details 2'}
        }
        Gen_Document(mp)
        self.assertTrue(os.path.exists('Project Status Week ' + str(datetime.now().isocalendar()[1]) + ' 2023.docx'))

    #  Tests that the header alignment is set correctly
    def test_header_alignment(self):
        mp = {
            'project name': {1: 'Project 1', 2: 'Project 2'},
            'Engineer name': {1: 'Engineer 1', 2: 'Engineer 2'},
            'Related To': {1: 'Related To 1', 2: 'Related To 2'},
            'status': {1: 'In Progress', 2: 'Completed'},
            'Details': {1: 'Details 1', 2: 'Details 2'}
        }
        Gen_Document(mp)
        doc = Document('Project Status Week ' + str(datetime.now().isocalendar()[1]) + ' 2023.docx')
        header = doc.sections[0].header
        self.assertEqual(header.paragraphs[0].alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)

    #  Tests that the current date is added to the header
    def test_add_current_date(self):
        mp = {
            'project name': {1: 'Project 1', 2: 'Project 2'},
            'Engineer name': {1: 'Engineer 1', 2: 'Engineer 2'},
            'Related To': {1: 'Related To 1', 2: 'Related To 2'},
            'status': {1: 'In Progress', 2: 'Completed'},
            'Details': {1: 'Details 1', 2: 'Details 2'}
        }
        Gen_Document(mp)
        doc = Document('Project Status Week ' + str(datetime.now().isocalendar()[1]) + ' 2023.docx')
        header = doc.sections[0].header
        self.assertEqual(header.paragraphs[1].text, datetime.now().strftime("%B %d, %Y"))

    #  Tests that rows are highlighted based on the status value
    def test_highlight_rows(self):
        mp = {
            'project name': {1: 'Project 1', 2: 'Project 2'},
            'Engineer name': {1: 'Engineer 1', 2: 'Engineer 2'},
            'Related To': {1: 'Related To 1', 2: 'Related To 2'},
            'status': {1: 'In Progress', 2: 'Completed'},
            'Details': {1: 'Details 1', 2: 'Details 2'}
        }
        Gen_Document(mp)
        doc = Document('Project Status Week ' + str(datetime.now().isocalendar()[1]) + ' 2023.docx')
        table = doc.tables[0]
        self.assertEqual(table.cell(1, 3).paragraphs[0].runs[0].font.color.rgb, RGBColor(0, 0, 255))
        self.assertEqual(table.cell(2, 3).paragraphs[0].runs[0].font.color.rgb, RGBColor(0, 255, 0))

    #  Tests that the function handles an empty mp
    def test_empty_mp(self):
        mp = {}
        Gen_Document(mp)
        self.assertTrue(os.path.exists('Project Status Week ' + str(datetime.now().isocalendar()[1]) + ' 2023.docx'))
