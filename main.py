import csv
import os
from datetime import datetime, timedelta
import pandas as pd
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from multiprocessing import context
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
from keras.src.applications.densenet import layers
from selenium import webdriver
from selenium.webdriver import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from docx import Document
from selenium.webdriver.edge.options import Options
import tensorflow as tf
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np

# Global variables
size = 33
GREEN = "\033[32m"
RESET = "\033[0m"
RED = "\033[31m"
PASS = 0
FAIL = 0
username = "Abdelrahman.Rasem"
password = "20111940@A"
os.environ['PATH'] += r"C:\Users\Abdelrahman.Rasem\Downloads\edgedriver_win64 (1)\msedgedriver.exe"
startDate = {}
actend = {}
end = {}
preduction_period = {}


def send_email(fn):
    """
    Sends an email with an attachment.

    Args:
        fn (str): The filename of the attachment to be sent.

    """
    try:
        # Create a message object
        msg = MIMEMultipart()

        # Get the current week of the year
        weekemail = str(datetime.now().isocalendar()[1])

        # Set email details (sender, recipient, subject)
        msg['From'] = 'globtesting0@gmail.com'
        msg['To'] = 'Mohammad.Farah@globitel.com'
        # msg['To'] = 'Abdelrahman.Rasem@globitel.com'
        # Add CC recipients
        msg['CC'] = 'Abdelrahman.Rasem@globitel.com'

        # Set the subject of the email
        msg['Subject'] = 'Project Status ' + weekemail + ' 2023'

        # Attach the file to the email
        filename2email = fn
        if not os.path.exists(fn):
            raise FileNotFoundError("Attachment file not found")
        attachment = open(filename2email, 'rb')

        # Create a MIMEBase object to represent the attachment
        part = MIMEBase('application', 'octet-stream')

        # Read the attachment file and encode it using base64
        part.set_payload(attachment.read())
        encoders.encode_base64(part)

        # Set the header for the attachment
        part.add_header('Content-Disposition', f'attachment; filename="{filename2email}"')

        # Attach the attachment to the email message
        msg.attach(part)

        # Connect to the email server and send the email
        server = smtplib.SMTP('smtp.gmail.com', 587)

        # Start a secure TLS connection
        server.starttls()

        # Login to the email account
        server.login('globtesting0@gmail.com', 'vfqmlzzwiubkyxlg')

        # Send the email
        server.send_message(msg)

        # Quit the email server connection
        server.quit()

    except Exception as error:
        # print("An error occurred:", error)
        print(type(error).__name__)
        exception = context.exception
        # Additional code to handle or raise the exception as needed
        raise exception
        # print("Error message:", str(error))


# Objective: The Gen_Document function aims to generate a Word document containing a table with project status
# information. The function takes a hash map as input, which contains project details such as project name,
# engineer name, related to, status, and comments. The function highlights the rows of the table based on the status
# value and saves the document with a filename that includes the current week number.
#
# Inputs:
# - mp: a hash map containing project details such as project name, engineer name, related to, status, and comments.
#
# Flow:
# 1. Create a new Word document.
# 2. Add a header to the document with the title "Globitel Project Status Report" and the current date.
# 3. Create a table with headers for project name, engineer name, related to, status, and comments.
# 4. Iterate over the hash map and add data to the table.
# 5. Highlight rows based on the status value.
# 6. Save the document with a filename that includes the current week number.
# 7. Modify the page size for landscape orientation.
# 8. Adjust the width of the last column in the table.
# 9. Save the modified document.
#
# Outputs:
# - A Word document containing a table with project status information.
# - The document is saved with a filename that includes the current week number.
#
# Additional aspects:
# - The function uses the docx library to create and modify Word documents.
# - The function uses RGBColor to set the font color of the status cell based on the status value.
# - The function assumes that the table is the first and only table in the document.
# - The function adjusts the width of the last column in the table to 4 inches.


def Data_Analysis(mp):
    # Convert the dictionary to a pandas DataFrame
    df = pd.DataFrame(mp)

    csv_filename = 'my_data '+str(datetime.now().isocalendar()[1])+' .csv'

    # Open the CSV file in write mode
    with open(csv_filename, 'w', newline='') as csv_file:
        csv_writer = csv.writer(csv_file)

        # Write the header row using the dictionary keys
        csv_writer.writerow(mp.keys())

        # Write the data as a single row
        csv_writer.writerow(mp.values())

    print(f'{csv_filename} has been created.')

    # Scatter plot with day against tip
    plt.plot(mp['status'])
    plt.plot(mp['project name'])

    # Adding Title to the Plot
    plt.title("Scatter Plot")

    # Setting the X and Y labels
    plt.xlabel('status')
    plt.ylabel('project name')
    plt.savefig('Project .png')




    # Perform data analysis or generate statistics using pandas methods
    summary_stats = df.describe()
    status_counts = df['status'].value_counts()
    average_duration = df['Duration'].mean()
    Eng_Count = df['Engineer name'].value_counts()

    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Engineer Statistics Report', level=1)

    # Create a table of contents
    doc.add_heading('Table of Contents', level=2)

    # Add a field instruction for the TOC
    toc_field_code = r'TOC \o "1-3" \h \z \u'
    toc_xml = r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wx="http://schemas.microsoft.com/office/word/2003/auxHint" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" w:rsidR="005A15A2" w:rsidRPr="005A15A2" w:rsidRDefault="005A15A2" w:rsidP="005A15A2"><w:pPr><w:pStyle w:val="TOCHeading"/><w:rPr><w:noProof/></w:rPr></w:pPr><w:r><w:t></w:t></w:r><w:r w:rsidR="005A15A2"><w:fldSimple w:instr="TOC \o &quot;1-3&quot; \h \z \u"><w:rPr><w:noProof/></w:rPr></w:fldSimple></w:r></w:p>'
    toc_paragraph = parse_xml(toc_xml)
    doc.element.body.insert(2, toc_paragraph)

    # Add summary statistics as a table to the document
    doc.add_heading('Summary Statistics', level=2)
    doc.add_paragraph(str(summary_stats))

    # Add statistics about each engineer
    doc.add_heading('Engineer Statistics', level=2)
    for engineer_name in df['Engineer name'].unique():
        doc.add_heading(f'Statistics for Engineer: {engineer_name}', level=3)
        engineer_df = df[df['Engineer name'] == engineer_name]
        engineer_stats = engineer_df.describe()
        doc.add_paragraph(str(engineer_stats))

    # Calculate overall statistics for all engineers
    overall_stats = df.describe()

    # Add overall minimum and maximum values to the document
    doc.add_heading('Overall Statistics', level=3)
    doc.add_paragraph(f'Minimum values:\n{overall_stats.loc["min"]}\n')
    doc.add_paragraph(f'Maximum values:\n{overall_stats.loc["max"]}\n')

    # Create a pie chart
    plt.figure(figsize=(8, 8))
    engineer_counts = Eng_Count
    engineer_counts.plot(kind='pie', autopct='%1.1f%%')
    plt.xlabel('Engineer Name')
    plt.ylabel('Count')
    plt.title('Engineer Distribution')
    plt.tight_layout()
    plt.savefig('Engineer_distribution.png')
    plt.close()

    # Add the pie chart to the document
    doc.add_picture('Engineer_distribution.png', width=Inches(6))

    # Add spacing after the pie chart
    doc.add_paragraph()

    plt.figure(figsize=(10, 12))
    Eng_Count.plot(kind='bar')
    plt.xlabel('ENG Name')
    plt.ylabel('Count')
    plt.title('ENG Counts')
    plt.tight_layout()
    plt.savefig('ENG_count.png')
    plt.close()

    # Add status counts chart to the Word document
    doc.add_heading('Engineer Project Load', level=2)
    doc.add_picture('ENG_count.png', width=Inches(6), height=Inches(8))

    # Generate and save status counts bar chart as an image
    plt.figure(figsize=(6, 4))
    status_counts.plot(kind='bar')
    plt.xlabel('Status')
    plt.ylabel('Count')
    plt.title('Status Counts')
    plt.tight_layout()
    plt.savefig('status_counts_chart.png')
    plt.close()

    # Add status counts chart to the Word document
    doc.add_heading('Status Counts', level=2)
    doc.add_picture('status_counts_chart.png', width=Inches(6), height=Inches(4))

    # Add average duration to the document
    doc.add_heading('Average Duration', level=2)
    formatted_average_duration = "{:.2f}".format(average_duration)
    doc.add_paragraph(f"The average duration is: {formatted_average_duration} days")

    # Save the document
    doc.save('data_analysis_report.docx')





# Data_Analysis(mp)
# Objective: - The objective of the Gen_Document function is to generate a Word document containing a table with
# project status information.
#
# Inputs:
# - mp: a dictionary containing project status information.
#
# Flow:
# - Create a new Word document.
# - Add a header to the document.
# - Add the current date to the header.
# - Add a paragraph with the total number of running projects.
# - Create a table with headers.
# - Iterate over the project status information and add data to the table.
# - Highlight rows based on the status value.
# - Modify the page size for landscape orientation.
# - Adjust the width of the last column in the table.
# - Save the modified document.
#
# Outputs:
# - A Word document containing a table with project status information.
#
# Additional aspects:
# - The function uses the python-docx library to create and modify Word documents.
# - The function highlights rows in the table based on the status value.
# - The function adjusts the page size and column width for better document formatting.
# - The function saves the modified document with a filename based on the current week number.
def Gen_Document(mp):
    if not mp:
        return
    # Create a new Word document
    doc = Document()

    # Add a header to the document
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Globitel Project Status Report"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_paragraph.style = doc.styles["Heading 1"]

    # Add the current date on the left side of the header
    date_paragraph = header.add_paragraph()
    date_paragraph.text = datetime.now().strftime("%B %d, %Y")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Add a new paragraph with a line break and additional text
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"Total Running Projects {size}")

    # Create a table with headers
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'project name'
    header_cells[1].text = 'Engineer name'
    header_cells[2].text = 'Related To'
    header_cells[3].text = 'status'
    header_cells[4].text = 'Comments'

    # Set the header cells' formatting
    for cell in header_cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    status_cell = table.cell(0, 4)
    status_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    status_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status_cell.paragraphs[0].runs[0].bold = True
    status_cell.paragraphs[0].runs[0].font.size = Pt(12)
    status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red color

    # Iterate over the hash map and add data to the table
    for number in range(1, size):
        row_cells = table.add_row().cells
        row_cells[0].text = mp['project name'].get(number,
                                                   '')  # Get project name, defaulting to empty string if not found
        row_cells[1].text = mp['Engineer name'].get(number,
                                                    '')  # Get engineer name, defaulting to empty string if not found
        row_cells[2].text = mp['Related To'].get(number,
                                                 '')  # Get engineer name, defaulting to empty string if not found
        row_cells[3].text = mp['status'].get(number, '')  # Get status, defaulting to empty string if not found
        row_cells[4].text = mp['Details'].get(number, '')

    # Highlight rows based on the status value
    for row in table.rows:
        status_cell = row.cells[3]
        status = status_cell.text.strip().lower()

        # Define the highlighting color based on the status value
        if status == 'in progress':
            highlight_color = RGBColor(0, 0, 255)  # Blue
        elif status == 'completed':
            highlight_color = RGBColor(0, 255, 0)  # Green
        elif status == 'on hold':
            highlight_color = RGBColor(255, 0, 0)  # Red
        elif status == 'Initiated':
            highlight_color = RGBColor(64, 64, 64)  # Gray
        else:
            highlight_color = RGBColor(64, 64, 64)  # Gray

        if highlight_color:
            status_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            status_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            status_cell.paragraphs[0].runs[0].bold = True
            status_cell.paragraphs[0].runs[0].font.size = Pt(12)
            status_cell.paragraphs[0].runs[0].font.color.rgb = highlight_color

    current_week = str(datetime.now().isocalendar()[1])
    filename = 'Project Status Week ' + current_week + ' 2023.docx'
    doc.save(filename)

    # Open the Word document
    doc = Document(filename)

    # Get the first section of the document
    section = doc.sections[0]

    # Modify the page size for landscape orientation
    section.page_width = Pt(792)  # Set the width to 11 inches (792 points)
    section.page_height = Pt(612)  # Set the height to 8.5 inches (612 points)

    # Assuming the table is the first and only table in the document
    table = doc.tables[0]

    # Get the last column index
    last_column_index = len(table.columns) - 1

    # Iterate over the cells in the last column and adjust the width
    for row in table.rows:
        cell = row.cells[last_column_index]
        cell.width = Inches(4)  # Adjust the width as needed

    # Save the modified document
    doc.save(filename)
    # Data_Analysis(mp)


# Objective: The objective of the function is to collect data from a web page, create a Word document, and highlight
# rows in the table based on the status value. The function also prints the number of passed and failed projects,
# generates a file name, and sends an email if all projects pass.
#
# Inputs:
# The function takes no inputs, but it uses global variables for the username, password, and size of the project.
#
# Flow:
# 1. The function creates an EdgeOptions object and a new instance of the Edge driver.
# 2. It navigates to the login page and enters the username and password.
# 3. The function opens a new tab and navigates to a new page.
# 4. It finds all buttons on the page by their tag name and switches to the second tab.
# 5. The function collects data from the web page and stores it in a hash map.
# 6. It highlights rows based on the status value and saves the Word document.
# 7. The function prints the number of passed and failed projects and sends an email if all projects pass.
# 8. Finally, the function closes the browser window.
#
# Outputs: The function generates a Word document with highlighted rows based on the status value. It also prints the
# number of passed and failed projects and sends an email if all projects pass.
#
# Additional aspects: The function uses the selenium, docx, and datetime libraries to interact with the web page,
# create a Word document, and get the current date. It also uses the Options class to add a headless argument and
# maximize the window. The function adds a delay of 3 seconds to avoid errors and scrolls the page to view the last
# three projects. The function uses global variables for the username, password, and size of the project and updates
# the size variable based on the number of projects on the web page.

def Collect_Data():
    global PASS, FAIL, username, password, GREEN, RED, RESET, size
    # Create EdgeOptions object
    options = Options()

    # Create a new instance of the Edge driver
    browser = webdriver.Edge(options=options)

    # Navigate to the login page
    browser.get("https://supportcrm.globitel.com/index.php")
    browser.maximize_window()

    # Find the username and password input fields and enter your credentials
    username_input = browser.find_element(By.ID, "username")
    username_input.send_keys(username)
    password_input = browser.find_element(By.ID, "password")
    password_input.send_keys(password)
    password_input.send_keys(Keys.RETURN)

    # select first tab to close it.
    first_tab_handle = browser.window_handles[0]

    # waiting for page to load.
    browser.implicitly_wait(3)

    browser.get(
        "https://supportcrm.globitel.com/index.php?module=Project&view=List&viewname=183&app=MARKETING")
    # Open a new tab and navigate to a new page
    browser.execute_script(
        "window.open('https://supportcrm.globitel.com/index.php?module=Project&view=List&viewname=183&app=MARKETING',"
        "'_blank')")

    actions = ActionChains(browser)
    actions.send_keys(Keys.END).perform()

    # listopen=browser.find_element(By.CLASS_NAME, "app-icon fa fa-bars")

    # Find all buttons on the page by their tag name
    secound_tab = browser.window_handles[1]

    browser.switch_to.window(first_tab_handle)
    browser.close()
    browser.switch_to.window(secound_tab)

    print("---------------------------------- target elem --------------------------------")

    mp = {'project name': {}, 'Engineer name': {}, 'Related To': {}, 'status': {}, 'Details': {}, 'Duration': {}}

    i = 9
    size = int(
        browser.find_element(By.XPATH, '//*[@id="listview-actions"]/div/div[3]/div/span/span[1]').text.split(' to ')[
            1]) + 1
    print("Starting Test on Total project size = ", size)
    number = 1  # Initialize the loop variable

    for number in range(1,size):
    #while number < size:
        time.sleep(3)  # Add a delay of 3 seconds
        mp['project name'][number] = browser.find_element(By.XPATH,
                                                          '//*[@id="Project_listView_row_' + str(
                                                              number) + '"]/td[2]').text

        if number == 23:
            print("DEBUG")

        if mp['project name'][number] =="Test":
            PASS=PASS+1
            #number+=1
            #i+=1
            continue
        try:
            mp['Engineer name'][number] = browser.find_element(By.XPATH,
                                                               '//*[@id="Project_listView_row_' + str(
                                                                   number) + '"]/td[4]').text
            if mp['Engineer name'][number] == "":
                print("Engineer name is empty")




            mp['Related To'][number] = browser.find_element(By.XPATH,
                                                            '//*[@id="Project_listView_row_' + str(
                                                                number) + '"]/td[7]/span[1]/span').text
            #update to 8 if working in allproject or 7


            if mp['Related To'][number] == "":
                print("Related To is empty")

            mp['status'][number] = browser.find_element(By.XPATH,
                                                        '//*[@id="Project_listView_row_' + str(
                                                            number) + '"]/td[6]').text

            if mp['status'][number] == "":
                print("status is empty")
                #//*[@id="mCSB_68_container"]/div[3]/table/tbody/tr[5]/td[2]/div/span/span
                #//*[@id="mCSB_66_container"]/div[3]/table/tbody/tr[5]/td[2]/div/span/span
                #//*[@id="mCSB_67_container"]/div[3]/table/tbody/tr[5]/td[2]/div/span/span
                # mp['status'][number] = browser.find_element(By.XPATH,
                #                                             '//*[@id="Project_listView_row_' + str(
                #                                                 number) + '"]/td[6]').text

            browser.find_element(By.XPATH, '//*[@id="Project_listView_row_' + str(
                number) + '"]/td[1]/div/span[2]/a').click()

        except Exception as e:
            print("project : ", mp['project name'][number], RED + "FAIL" + RESET + str(e))
            FAIL = FAIL + 1
            continue

        time.sleep(3)  # Add a delay of 3 seconds
        try:
            mp['Details'][number] = browser.find_element(By.XPATH,
                                                         '//*[@id="detailView"]/div/div/div/div[2]/div/div['
                                                         '3]/div/div[1]').text
            if mp['Details'][number] == "":
                print("Details is empty")
#//*[@id="mCSB_32_container"]/div[3]/table/tbody/tr[2]/td[2]/div/span
            # from_Date = browser.find_element(By.XPATH, '//*[@id="mCSB_' + str(
            #     i) + '_container"]/div[3]/table/tbody/tr[2]/td[2]/div/span').text
            # to_Date = browser.find_element(By.XPATH, '//*[@id="mCSB_' + str(
            #     i) + '_container"]/div[3]/table/tbody/tr[4]/td[2]/div/span').text
                #//*[@id="Project_listView_row_32"]/td[3]/span[1]/span
            from_Date = browser.find_element(By.XPATH, '//*[@id="Project_listView_row_'+str(number)+'"]/td[3]/span[1]/span').text
            to_Date = browser.find_element(By.XPATH, '//*[@id="Project_listView_row_'+str(number)+'"]/td[5]/span[1]/span').text
            date_format = "%Y-%m-%d"
            # Get the current date
            current_date = datetime.now().date()

            if to_Date == '':
                to_Date = current_date + timedelta(days=90)
                to_Date = to_date.strftime("%Y-%m-%d")

            from_date = datetime.strptime(from_Date, date_format)
            to_date = datetime.strptime(to_Date, date_format)
            mp['Duration'][number] = to_date - from_date

            browser.find_element(By.XPATH,
                                 '//*[@id="mCSB_' + str(i) + '_container"]/div[1]/div[2]/button').click()
        except Exception as e:
            print("project : ", mp['project name'][number], RED + "FAIL" + RESET)
            print(f"Fail Reason : {e}")
            FAIL = FAIL + 1
            i+=1
            continue

        i = i + 1
        print("project : ", mp['project name'][number], GREEN + "PASS" + RESET+" Progress: "+str((number/size)*100))
        PASS = PASS + 1

        if number > 3:
            browser.execute_script("arguments[0].scrollIntoView(true);",
                                   browser.find_element(By.XPATH, '//*['
                                                                  '@id="Project_listView_row_' + str(
                                       number - 3) + '"]/td[1]'))
        # try:
        #     if number==size-1:
        #         browser.execute_script("window.scrollTo(0, 0);")
        #         next_btn= browser.find_element(By.XPATH,'//*[@id="NextPageButton"]')#.click()
        #         if next_btn.is_enabled():
        #             next_btn.click()
        #             print("next page")
        #         else:
        #             print("end of content")
        #         size = int(
        #             browser.find_element(By.XPATH, '//*[@id="listview-actions"]/div/div[3]/div/span/span[1]').text.split(
        #                 ' to ')[
        #                 1]) + 1
        # except:
        #     continue

        #number+=1



    print(GREEN + "PASS: " + RESET, PASS, "/", size - 1)
    print(RED + "FAIL: " + RESET, FAIL, "/", size - 1)

    Gen_Document(mp)
    current_week = str(datetime.now().isocalendar()[1])
    filename = 'Project Status Week ' + current_week + ' 2023.docx'
    if PASS == size - 1:
        print("Sending Email....")
        # uncomment the next line to send email
        #send_email(filename)
    # #Schedule the email to be sent every day at 9:00 AM
    # schedule.every().week.monday.at('01:00').do(lambda: send_email("Project Status Week 24 2023.docx"))
    #
    # while True:
    #     schedule.run_pending()
    #     time.sleep(1)

    # Close the browser window
    Data_Analysis(mp)
    browser.quit()


def data_train():
    global startDate, actend, end, preduction_period
    # Create EdgeOptions object
    options = Options()

    # Create a new instance of the Edge driver
    browser = webdriver.Edge(options=options)

    # Navigate to the login page
    browser.get("https://supportcrm.globitel.com/index.php")
    browser.maximize_window()

    # Find the username and password input fields and enter your credentials
    username_input = browser.find_element(By.ID, "username")
    username_input.send_keys(username)
    password_input = browser.find_element(By.ID, "password")
    password_input.send_keys(password)
    password_input.send_keys(Keys.RETURN)

    # select first tab to close it.
    first_tab_handle = browser.window_handles[0]

    # waiting for page to load.
    browser.implicitly_wait(3)
    browser.get(
        "https://supportcrm.globitel.com//index.php?module=Project&view=List&viewname=213&app=MARKETING")
    # Open a new tab and navigate to a new page
    browser.execute_script(
        "window.open('https://supportcrm.globitel.com//index.php?module=Project&view=List&viewname=213&app=MARKETING')")

    actions = ActionChains(browser)
    actions.send_keys(Keys.END).perform()

    # listopen=browser.find_element(By.CLASS_NAME, "app-icon fa fa-bars")

    # Find all buttons on the page by their tag name
    secound_tab = browser.window_handles[1]

    browser.switch_to.window(first_tab_handle)
    browser.close()
    browser.switch_to.window(secound_tab)

    date_format = "%Y-%m-%d"
    # Get the current date

    # '//*[@id="Project_listView_row_1"]/td[5]/span[1]/span'
    for num in range(1, 11):
        startDate[num] = browser.find_element(By.XPATH, '//*[@id="Project_listView_row_' + str(
            num) + '"]/td[5]/span[1]/span').text
        actend[num] = browser.find_element(By.XPATH,
                                           '//*[@id="Project_listView_row_' + str(num) + '"]/td[6]/span[1]/span').text
        end[num] = browser.find_element(By.XPATH,
                                        '//*[@id="Project_listView_row_' + str(num) + '"]/td[7]/span[1]/span').text
        from_date = datetime.strptime(startDate[num], date_format)
        to_date = datetime.strptime(end[num], date_format)
        preduction_period[num] = to_date - from_date
        from_date = datetime.strptime(startDate[num], date_format)
        to_date = datetime.strptime(actend[num], date_format)
        actend[num] = to_date - from_date

    print("pass")


# send_email('test.txt')
# {1: datetime.timedelta(days=56), 2: datetime.timedelta(days=56), 3: datetime.timedelta(days=56), 4: datetime.timedelta(days=42), 5: datetime.timedelta(days=56), 6: datetime.timedelta(days=56), 7: datetime.timedelta(days=39), 8: datetime.timedelta(days=175), 9: datetime.timedelta(days=30), 10: datetime.timedelta(days=35)}
# {1: datetime.timedelta(days=355), 2: datetime.timedelta(days=85), 3: datetime.timedelta(days=49), 4: datetime.timedelta(days=9), 5: datetime.timedelta(days=27), 6: datetime.timedelta(days=18), 7: datetime.timedelta(days=123), 8: datetime.timedelta(days=652), 9: datetime.timedelta(days=23), 10: datetime.timedelta(days=15)}


Collect_Data()
# Data_Analysis()


# # Prepare the training data
# x_train = np.array([56, 56, 56, 42, 56, 56, 39, 175, 30, 35])  # Estimated project days
# y_train = np.array([355, 85, 49, 9, 27, 18, 123, 652, 23, 15])  # Actual project days
#
# # Define the model architecture
# model = tf.keras.Sequential([
#     layers.Dense(64, activation='relu', input_shape=(1,)),
#     layers.Dense(64, activation='relu'),
#     layers.Dense(1)
# ])
#
# # Compile the model
# model.compile(optimizer='adam',
#               loss='mean_squared_error',
#               metrics=['mae'])
#
# # Train the model
# model.fit(x_train, y_train, epochs=10, batch_size=32)
#
# # Predict on new data
# x_test = np.array([79, 107, 59, 19, 46, 56, 42])  # Estimated project days for new data
# y_pred = model.predict(x_test)
# print(y_pred)


# num_classes=4
# input_dim=4
# # Define the model architecture
#
# import numpy as np
#
# # Define the model architecture
# model = tf.keras.Sequential([
#     layers.Dense(64, activation='relu', input_shape=(1,)),
#     layers.Dense(64, activation='relu'),
#     layers.Dense(1)
# ])
#
# # Compile the model
# model.compile(optimizer='adam',
#               loss='mean_squared_error',
#               metrics=['mae'])
#
# # Prepare the training and validation data
# x_train = np.array([56, 56, 56, 42, 56, 56, 39, 175, 30, 35])
# y_train = np.array([355, 85, 49, 9, 27, 18, 123, 652, 23, 15])
# x_val = np.array([56, 56, 56, 42, 56, 56, 39, 175, 30, 35])
# y_val = np.array([355, 85, 49, 9, 27, 18, 123, 652, 23, 15])
#
# # Train the model
# model.fit(x_train, y_train, epochs=10, batch_size=32, validation_data=(x_val, y_val))
#
#
# x_test={1: 42, 2: 79, 3: 107,
#                    4: 59, 5: 19, 6: 46,
#                    7: 56, 8: 42, 9: 21,
#                    10: 29, 11: 87, 12: 76,
#                    13: 61, 14: 84, 15: 101,
#                    16: 59, 17: 72, 18: 306,
#                    19: 382, 20: 22, 21: 92,
#                    22: 84, 23: 61, 24: 40,
#                    25: 56, 26: 153, 27: 81,
#                    28: 118, 29: 92, 30: 92,
#                    31: 59}
# # Convert the test data to a numpy array
# x_test = np.array(list(x_test.values()))
# # Evaluate the model
# test_loss, test_acc = model.evaluate(x_test)
# print('Test accuracy:', test_acc)
# print('Test loss:', test_loss)
#
